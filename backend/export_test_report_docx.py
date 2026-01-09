"""
Convert Comprehensive Test Report from Markdown to DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_markdown_to_docx(md_file, docx_file, screenshots_dir):
    """Convert markdown file to DOCX with formatting"""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (# )
        if line.startswith('# '):
            title = line[2:]
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Heading 2 (## )
        elif line.startswith('## '):
            heading = line[3:]
            doc.add_heading(heading, level=1)
            
        # Heading 3 (### )
        elif line.startswith('### '):
            heading = line[4:]
            doc.add_heading(heading, level=2)
            
        # Heading 4 (#### )
        elif line.startswith('#### '):
            heading = line[5:]
            doc.add_heading(heading, level=3)
            
        # Heading 5 (##### )
        elif line.startswith('##### '):
            heading = line[6:]
            doc.add_heading(heading, level=4)
            
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph()
            
        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Handle checkboxes
            text = text.replace('- [ ]', '☐').replace('- [x]', '☑').replace('- [/]', '◐')
            text = text.replace('✅', '✓').replace('⚠️', '⚠').replace('❌', '✗')
            doc.add_paragraph(text, style='List Bullet')
            
        # Code blocks
        elif line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)
            # Add shading
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p._p.get_or_add_pPr().append(shading_elm)
            
        # Images
        elif line.startswith('![') and '](' in line:
            # Extract image path
            match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
            if match:
                caption = match.group(1)
                img_path = match.group(2)
                
                # Convert file:/// URLs to local paths
                if img_path.startswith('file:///'):
                    img_path = img_path[8:]  # Remove file:///
                    img_path = img_path.replace('/', '\\')
                
                # Add caption
                if caption:
                    p = doc.add_paragraph()
                    run = p.add_run(caption)
                    run.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image if file exists
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        doc.add_paragraph(f"[Image not found: {img_path}]")
                else:
                    doc.add_paragraph(f"[Image not found: {img_path}]")
        
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Parse table
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip())
                i += 1
            i -= 1
            
            # Skip separator line if present
            if len(table_lines) > 1 and re.match(r'\|[\s\-:|]+\|', table_lines[1]):
                table_lines.pop(1)
            
            # Create table
            rows = []
            for tline in table_lines:
                cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                rows.append(cells)
            
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        
                        # Bold header row
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        
        # Regular paragraph
        else:
            text = line
            # Handle inline formatting
            text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*([^\*]+)\*', r'\1', text)  # Italic
            text = re.sub(r'`([^`]+)`', r'\1', text)  # Code
            
            # Handle special characters
            text = text.replace('✅', '✓').replace('⚠️', '⚠').replace('❌', '✗')
            
            p = doc.add_paragraph(text)
            
            # Apply formatting
            if '**' in line:
                matches = re.finditer(r'\*\*([^\*]+)\*\*', line)
                for match in matches:
                    for run in p.runs:
                        if match.group(1) in run.text:
                            run.bold = True
            
            if '✅' in line or 'PASSED' in line:
                for run in p.runs:
                    if '✓' in run.text or 'PASSED' in run.text:
                        run.font.color.rgb = RGBColor(0, 128, 0)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Document saved to: {docx_file}")

if __name__ == "__main__":
    # Paths
    artifacts_dir = r"C:\Users\hp\.gemini\antigravity\brain\c064c695-5740-4d63-a5a1-17ca1bb02b9d"
    md_file = os.path.join(artifacts_dir, "comprehensive_test_report.md")
    docx_file = os.path.join(artifacts_dir, "BioAttend_Comprehensive_Test_Report.docx")
    screenshots_dir = artifacts_dir
    
    print("Converting markdown to DOCX...")
    print(f"Input: {md_file}")
    print(f"Output: {docx_file}")
    
    parse_markdown_to_docx(md_file, docx_file, screenshots_dir)
    
    print("\n✓ Conversion complete!")
    print(f"\nExported file: {docx_file}")
